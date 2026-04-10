# -*- coding: utf-8 -*-
"""生成「仿真数据集制作」论文章节 .docx。运行: D:\\Anaconda\\envs\\rocket\\python.exe tools/generate_simulation_dataset_chapter_docx.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "论文章节_仿真数据集制作.docx"

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("仿真数据集构建方法")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Times New Roman"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.paragraph_format.space_after = Pt(18)

    add_para(
        doc,
        "为在可控条件下验证可见光着陆标志关键点检测与位姿估计算法，本文构建物理轨迹仿真与三维渲染相结合的合成数据集。"
        "六自由度轨迹由独立 Python 程序数值生成；三维场景、传感器与图像序列在 Blender 中依据该轨迹驱动得到。"
        "数据与标注文件存放于 datasets 目录，其中渲染结果、JSON 标注与预训练权重等路径通常列入版本管理忽略规则，以本地生成物为准。"
        "下文分节说明标志物设计、轨迹生成、场景搭建、数据采集、标注方案及数据集特性。",
    )

    add_heading(doc, "1  着陆标志物设计", level=1)
    add_para(
        doc,
        "着陆区采用圆形标志与高对比度 H 字形图案相结合的平面合作目标，布置于着陆平面（海拔 0）附近，以利于远距离检测与亚像素级角点定位。"
        "世界坐标系取 X 东、Y 北、Z 上；标志中心与着陆区几何中心与轨迹终值一致。"
        "在该平面上选取 9 个语义关键点：1 个几何中心、4 个主结构臂端点及 4 个内侧辅助点，共面于 Z=0，便于透视投影与 PnP 位姿解算。",
    )
    add_para(
        doc,
        "各点三维坐标由数据集配置文件 kp.json 给出（与项目说明文档中的关键点表一致）。"
        "参考尺度（单位 m）：中心 kp0 为 (0, 0, 0)；±X、±Y 臂端 kp1–kp4 距中心 21.5 m；内侧点 kp5–kp8 为 (±3.9, ±9.0, 0)。"
        "定稿时应与本地 kp.json 核对是否一致。",
    )

    add_heading(doc, "2  仿真轨迹生成", level=1)
    add_para(
        doc,
        "六自由度轨迹由 Simulation 目录下 generate_trajectory.py 在常规 Python 环境中数值积分得到。"
        "模型参考 Falcon 9 一级着陆段量级，包含变密度大气、气动阻力、推力可调发动机、燃料消耗及简化制导律，"
        "输出位置、欧拉姿态、速度与加速度时间序列，采样频率 50 Hz，存为 CSV，列包括 time, x, y, z, roll, pitch, yaw, vx, vy, vz, ax, ay, az，姿态角为弧度。",
    )
    add_para(
        doc,
        "可选地，采用 generate_vertical_recovery_traj.py 在保持位置与线运动不变的前提下，对 roll、pitch、yaw 叠加受高度包络调制的扰动，"
        "并在近地面衰减至零，以模拟回收段小幅姿态波动且避免触地帧非物理大角偏差。"
        "生成文件需与 Blender 脚本中 CSV_PATH 配置一致，以保证关键帧与仿真时间对齐。",
    )

    add_heading(doc, "3  仿真场景生成", level=1)
    add_para(
        doc,
        "三维场景与渲染管线在 Blender 5.0 中实现，脚本为 Simulation/rocket_trajectory_blender.py，须在 Blender Scripting 工作区内运行而非系统解释器。"
        "脚本读取轨迹 CSV，按帧率 50 Hz 为箭体插入位置与姿态关键帧；机载相机固连于箭体，局部安装姿态使视线朝向下视着陆区（脚本内给定欧拉角与杆臂位置）。"
        "地面采用外部沙漠环境资产，着陆区包含上述标志几何；渲染采用 Cycles，分辨率 1080×720 像素，设定采样数与降噪，并开启运动模糊以贴近动态成像。",
    )
    add_para(
        doc,
        "使用前需根据本机路径修改脚本中的 BASE_DIR、CSV_PATH 及环境资源 ENV_DIR。"
        "输出图像序列默认写入 Blender 工程相对路径下的 rocket_render 目录，可再整理至 datasets/rocket_render_01、rocket_render_02 等子目录供训练读取。",
    )

    add_heading(doc, "4  仿真数据采集", level=1)
    add_para(
        doc,
        "数据采集即为 Blender 批量渲染得到的 RGB 图像序列，帧序与轨迹 CSV 行序一一对应，便于关联真值位姿或生成稠密监督。"
        "图像幅面为 1080×720；在针孔模型与方形像素假设下，可与项目说明中的焦距—传感器几何关系一致地换算内参（例如基于 24 mm 焦距与 36 mm 传感器宽度得到 fx = fy = 720 像素，主点近似为图像中心）。",
    )
    add_para(
        doc,
        "仿真未引入镜头畸变，便于与理想针孔模型及 PnP 管线对接。"
        "各条轨迹的帧序列可独立存放于 datasets 下不同子目录，并与 sequences 类配置（若使用）一致编号。",
    )

    add_heading(doc, "5  数据集标注", level=1)
    add_para(
        doc,
        "关键点二维标注由三维—二维投影得到：依据相机内参、外参及 kp.json 中世界坐标计算像素坐标，并导出为 COCO 关键点格式 JSON（如 annotations_coco_keypoints.json 及按 train/val/test 划分的文件），含可见性字段。"
        "标注生成可由 datasets 目录下的 generate_coco_keypoints.py 等脚本复现，以保证与渲染几何一致。",
    )
    add_para(
        doc,
        "训练、验证与测试集可按约 7:2:1 划分；测试集可附加距离或高度分层标签（如 range_label），用于分段评估不同观测距离下的算法性能。",
    )

    add_heading(doc, "6  数据集特性", level=1)
    add_para(
        doc,
        "在典型实验配置下，采用两条独立下降轨迹，每条约 2000 帧，合计约 4000 帧，覆盖从较高海拔至近地面的视角与尺度变化。",
    )
    add_para(
        doc,
        "合成数据具有真值完备、标注与几何一致、可重复生成等优点；同时与真实传感器数据存在域差异，可通过域随机化、噪声叠加或实场数据微调加以缓解。"
        "建议在论文中用表格汇总轨迹条数、总帧数、划分比例、高度或斜距分布及关键点可见率，并与评测章节中的分层指标相对应。",
    )

    doc.save(out)
    print(f"已写入: {out}")


if __name__ == "__main__":
    main()
